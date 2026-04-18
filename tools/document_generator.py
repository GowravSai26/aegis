from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime


def generate_docx(state: dict) -> str:
    doc = Document()

    # Title
    title = doc.add_heading("CHARGEBACK DISPUTE RESPONSE", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    order = state.get("evidence_collected", {}).get("order_details", {})
    rows = [
        ("Case Reference", state.get("chargeback_id", "")),
        ("Merchant", order.get("merchant_name", "Unknown")),
        ("Transaction Date", str(order.get("order_date", ""))[:10]),
        ("Dispute Amount", f"${state.get('amount', 0):.2f}"),
        ("Reason Code", f"{state.get('reason_code')} — {state.get('reason_description')}"),
        ("Response Deadline", state.get("dispute_deadline", "")),
    ]
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Draft content
    draft = state.get("dispute_response_draft", "")

    # Strip the header we already rendered, write remaining sections
    sections_start = draft.find("EXECUTIVE SUMMARY")
    if sections_start != -1:
        body = draft[sections_start:]
    else:
        body = draft

    for line in body.splitlines():
        if line.strip() in (
            "EXECUTIVE SUMMARY", "EVIDENCE SUMMARY",
            "DETAILED ARGUMENT", "EVIDENCE ATTACHMENTS", "MERCHANT REQUEST"
        ):
            doc.add_heading(line.strip(), level=2)
        elif line.startswith("-" * 5):
            continue
        else:
            doc.add_paragraph(line)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Submitted by AEGIS Autonomous Dispute System\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    out_dir = Path("generated_docs")
    out_dir.mkdir(exist_ok=True)
    path = out_dir / f"{state.get('chargeback_id', 'dispute')}.docx"
    doc.save(str(path))
    return str(path)
